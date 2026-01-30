#!/usr/bin/env python3
"""
Wade-Giles to Pinyin Converter

This script reads a .docx or .pdf file, detects Wade-Giles romanization,
and converts it to toneless Pinyin based on the standard conversion table.

Usage:
    python wg_to_pinyin.py input.docx output.docx
    python wg_to_pinyin.py input.docx  # outputs to input_pinyin.docx
    python wg_to_pinyin.py input.pdf   # outputs to input_pinyin.pdf
"""

import re
import sys
import shutil
import zipfile
from pathlib import Path
from docx import Document

# Try to import PyMuPDF for PDF support
try:
    import fitz  # PyMuPDF
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False


# Wade-Giles to Pinyin conversion dictionary
# Based on "PINYIN CONVERSION SPECIFICATIONS - DICTIONARY STD: Standard pinyin conversion table"
#
# Notes on Wade-Giles orthography:
# - Apostrophe (') indicates aspiration: ch' -> ch, k' -> k, p' -> p, t' -> t, ts' -> c
# - Without apostrophe: ch -> zh, k -> g, p -> b, t -> d, ts -> z
# - ü is used after certain consonants (ch', hs, l, n, y)
# - Some syllables have variant spellings with diacritics (circumflex, breve) in original sources

# Core conversion dictionary (lowercase)
# Format: wade_giles -> pinyin
WG_TO_PINYIN = {
    # A
    'a': 'a',
    'ai': 'ai',
    'an': 'an',
    'ang': 'ang',
    'ao': 'ao',

    # CH (unaspirated -> zh)
    'cha': 'zha',
    'chai': 'zhai',
    'chan': 'zhan',
    'chang': 'zhang',
    'chao': 'zhao',
    'che': 'zhe',
    'chen': 'zhen',
    'cheng': 'zheng',
    'chi': 'ji',
    'chia': 'jia',
    'chiang': 'jiang',
    'chiao': 'jiao',
    'chieh': 'jie',
    'chien': 'jian',
    'chih': 'zhi',
    'chin': 'jin',
    'ching': 'jing',
    'chiu': 'jiu',
    'chiung': 'jiong',
    'cho': 'zhuo',
    'chou': 'zhou',
    'chu': 'zhu',
    'chua': 'zhua',
    'chuai': 'zhuai',
    'chuan': 'zhuan',
    'chuang': 'zhuang',
    'chui': 'zhui',
    'chun': 'zhun',
    'chung': 'zhong',
    'chü': 'ju',
    'chüan': 'juan',
    'chüeh': 'jue',
    'chün': 'jun',

    # CH' (aspirated -> ch/q)
    "ch'a": 'cha',
    "ch'ai": 'chai',
    "ch'an": 'chan',
    "ch'ang": 'chang',
    "ch'ao": 'chao',
    "ch'e": 'che',
    "ch'en": 'chen',
    "ch'eng": 'cheng',
    "ch'i": 'qi',
    "ch'ia": 'qia',
    "ch'iang": 'qiang',
    "ch'iao": 'qiao',
    "ch'ieh": 'qie',
    "ch'ien": 'qian',
    "ch'ih": 'chi',
    "ch'in": 'qin',
    "ch'ing": 'qing',
    "ch'iu": 'qiu',
    "ch'iung": 'qiong',
    "ch'o": 'chuo',
    "ch'ou": 'chou',
    "ch'u": 'chu',
    "ch'uai": 'chuai',
    "ch'uan": 'chuan',
    "ch'uang": 'chuang',
    "ch'ui": 'chui',
    "ch'un": 'chun',
    "ch'ung": 'chong',
    "ch'ü": 'qu',
    "ch'üan": 'quan',
    "ch'üeh": 'que',
    "ch'ün": 'qun',

    # E
    'en': 'en',
    'erh': 'er',
    'er': 'er',  # variant

    # F
    'fa': 'fa',
    'fan': 'fan',
    'fang': 'fang',
    'fei': 'fei',
    'fen': 'fen',
    'feng': 'feng',
    'fo': 'fo',
    'fou': 'fou',
    'fu': 'fu',

    # H
    'ha': 'ha',
    'hai': 'hai',
    'han': 'han',
    'hang': 'hang',
    'hao': 'hao',
    'hei': 'hei',
    'hen': 'hen',
    'heng': 'heng',
    'ho': 'he',
    'hou': 'hou',
    'hu': 'hu',
    'hua': 'hua',
    'huai': 'huai',
    'huan': 'huan',
    'huang': 'huang',
    'hui': 'hui',
    'hun': 'hun',
    'hung': 'hong',
    'huo': 'huo',

    # HS -> x
    'hsi': 'xi',
    'hsia': 'xia',
    'hsiang': 'xiang',
    'hsiao': 'xiao',
    'hsieh': 'xie',
    'hsien': 'xian',
    'hsin': 'xin',
    'hsing': 'xing',
    'hsiu': 'xiu',
    'hsiung': 'xiong',
    'hsü': 'xu',
    'hsüan': 'xuan',
    'hsüeh': 'xue',
    'hsün': 'xun',

    # I
    'i': 'yi',

    # J -> r
    'jan': 'ran',
    'jang': 'rang',
    'jao': 'rao',
    'je': 're',
    'jen': 'ren',
    'jeng': 'reng',
    'jih': 'ri',
    'jo': 'ruo',
    'jou': 'rou',
    'ju': 'ru',
    'juan': 'ruan',
    'jui': 'rui',
    'jun': 'run',
    'jung': 'rong',

    # K (unaspirated -> g)
    'ka': 'ga',
    'kai': 'gai',
    'kan': 'gan',
    'kang': 'gang',
    'kao': 'gao',
    'kei': 'gei',
    'ken': 'gen',
    'keng': 'geng',
    'ko': 'ge',
    'kou': 'gou',
    'ku': 'gu',
    'kua': 'gua',
    'kuai': 'guai',
    'kuan': 'guan',
    'kuang': 'guang',
    'kuei': 'gui',
    'kun': 'gun',
    'kung': 'gong',
    'kuo': 'guo',

    # K' (aspirated -> k)
    "k'a": 'ka',
    "k'ai": 'kai',
    "k'an": 'kan',
    "k'ang": 'kang',
    "k'ao": 'kao',
    "k'en": 'ken',
    "k'eng": 'keng',
    "k'o": 'ke',
    "k'ou": 'kou',
    "k'u": 'ku',
    "k'ua": 'kua',
    "k'uai": 'kuai',
    "k'uan": 'kuan',
    "k'uang": 'kuang',
    "k'uei": 'kui',
    "k'un": 'kun',
    "k'ung": 'kong',
    "k'uo": 'kuo',

    # L
    'la': 'la',
    'lai': 'lai',
    'lan': 'lan',
    'lang': 'lang',
    'lao': 'lao',
    'le': 'le',
    'lei': 'lei',
    'leng': 'leng',
    'li': 'li',
    'liang': 'liang',
    'liao': 'liao',
    'lieh': 'lie',
    'lien': 'lian',
    'lin': 'lin',
    'ling': 'ling',
    'liu': 'liu',
    'lo': 'luo',
    'lou': 'lou',
    'lu': 'lu',
    'luan': 'luan',
    'lun': 'lun',
    'lung': 'long',
    'lü': 'lü',
    'lüan': 'luan',
    'lüeh': 'lue',

    # M
    'ma': 'ma',
    'mai': 'mai',
    'man': 'man',
    'mang': 'mang',
    'mao': 'mao',
    'mei': 'mei',
    'men': 'men',
    'meng': 'meng',
    'mi': 'mi',
    'miao': 'miao',
    'mieh': 'mie',
    'mien': 'mian',
    'min': 'min',
    'ming': 'ming',
    'miu': 'miu',
    'mo': 'mo',
    'mou': 'mou',
    'mu': 'mu',

    # N
    'na': 'na',
    'nai': 'nai',
    'nan': 'nan',
    'nang': 'nang',
    'nao': 'nao',
    'nei': 'nei',
    'nen': 'nen',
    'neng': 'neng',
    'ni': 'ni',
    'niang': 'niang',
    'niao': 'niao',
    'nieh': 'nie',
    'nien': 'nian',
    'nin': 'nin',
    'ning': 'ning',
    'niu': 'niu',
    'no': 'nuo',
    'nu': 'nu',
    'nuan': 'nuan',
    'nung': 'nong',
    'nü': 'nü',
    'nüeh': 'nue',

    # O
    'o': 'e',
    'ou': 'ou',

    # P (unaspirated -> b)
    'pa': 'ba',
    'pai': 'bai',
    'pan': 'ban',
    'pang': 'bang',
    'pao': 'bao',
    'pei': 'bei',
    'pen': 'ben',
    'peng': 'beng',
    'pi': 'bi',
    'piao': 'biao',
    'pieh': 'bie',
    'pien': 'bian',
    'pin': 'bin',
    'ping': 'bing',
    'po': 'bo',
    'pu': 'bu',

    # P' (aspirated -> p)
    "p'a": 'pa',
    "p'ai": 'pai',
    "p'an": 'pan',
    "p'ang": 'pang',
    "p'ao": 'pao',
    "p'ei": 'pei',
    "p'en": 'pen',
    "p'eng": 'peng',
    "p'i": 'pi',
    "p'iao": 'piao',
    "p'ieh": 'pie',
    "p'ien": 'pian',
    "p'in": 'pin',
    "p'ing": 'ping',
    "p'o": 'po',
    "p'ou": 'pou',
    "p'u": 'pu',

    # S
    'sa': 'sa',
    'sai': 'sai',
    'san': 'san',
    'sang': 'sang',
    'sao': 'sao',
    'se': 'se',
    'sen': 'sen',
    'seng': 'seng',
    'sha': 'sha',
    'shai': 'shai',
    'shan': 'shan',
    'shang': 'shang',
    'shao': 'shao',
    'she': 'she',
    'shen': 'shen',
    'sheng': 'sheng',
    'shih': 'shi',
    'shou': 'shou',
    'shu': 'shu',
    'shua': 'shua',
    'shuai': 'shuai',
    'shuan': 'shuan',
    'shuang': 'shuang',
    'shui': 'shui',
    'shun': 'shun',
    'shuo': 'shuo',
    'so': 'suo',
    'sou': 'sou',
    'ssu': 'si',
    'su': 'su',
    'suan': 'suan',
    'sui': 'sui',
    'sun': 'sun',
    'sung': 'song',
    'szu': 'si',

    # T (unaspirated -> d)
    'ta': 'da',
    'tai': 'dai',
    'tan': 'dan',
    'tang': 'dang',
    'tao': 'dao',
    'te': 'de',
    'teng': 'deng',
    'ti': 'di',
    'tiao': 'diao',
    'tieh': 'die',
    'tien': 'dian',
    'ting': 'ding',
    'tiu': 'diu',
    'to': 'duo',
    'tou': 'dou',
    'tu': 'du',
    'tuan': 'duan',
    'tui': 'dui',
    'tun': 'dun',
    'tung': 'dong',

    # T' (aspirated -> t)
    "t'a": 'ta',
    "t'ai": 'tai',
    "t'an": 'tan',
    "t'ang": 'tang',
    "t'ao": 'tao',
    "t'e": 'te',
    "t'eng": 'teng',
    "t'i": 'ti',
    "t'iao": 'tiao',
    "t'ieh": 'tie',
    "t'ien": 'tian',
    "t'ing": 'ting',
    "t'o": 'tuo',
    "t'ou": 'tou',
    "t'u": 'tu',
    "t'uan": 'tuan',
    "t'ui": 'tui',
    "t'un": 'tun',
    "t'ung": 'tong',

    # TS (unaspirated -> z)
    'tsa': 'za',
    'tsai': 'zai',
    'tsan': 'zan',
    'tsang': 'zang',
    'tsao': 'zao',
    'tse': 'ze',
    'tsei': 'zei',
    'tsen': 'zen',
    'tseng': 'zeng',
    'tso': 'zuo',
    'tsou': 'zou',
    'tsu': 'zu',
    'tsuan': 'zuan',
    'tsui': 'zui',
    'tsun': 'zun',
    'tsung': 'zong',
    'tzu': 'zi',

    # TS' (aspirated -> c)
    "ts'a": 'ca',
    "ts'ai": 'cai',
    "ts'an": 'can',
    "ts'ang": 'cang',
    "ts'ao": 'cao',
    "ts'e": 'ce',
    "ts'en": 'cen',
    "ts'eng": 'ceng',
    "ts'o": 'cuo',
    "ts'ou": 'cou',
    "ts'u": 'cu',
    "ts'uan": 'cuan',
    "ts'ui": 'cui',
    "ts'un": 'cun',
    "ts'ung": 'cong',
    "tz'u": 'ci',

    # W
    'wa': 'wa',
    'wai': 'wai',
    'wan': 'wan',
    'wang': 'wang',
    'wei': 'wei',
    'wen': 'wen',
    'weng': 'weng',
    'wo': 'wo',
    'wu': 'wu',

    # Y
    'ya': 'ya',
    'yai': 'yai',
    'yang': 'yang',
    'yao': 'yao',
    'yeh': 'ye',
    'yen': 'yan',
    'yin': 'yin',
    'ying': 'ying',
    'yo': 'yo',
    'yu': 'you',
    'yung': 'yong',
    'yü': 'yu',
    'yüan': 'yuan',
    'yüeh': 'yue',
    'yün': 'yun',
}

# Build additional variants for ü spelled as 'u' after certain consonants
# In some sources, ü is written as plain 'u' after j, q, x, y (in WG: ch', hs, y)
# We already handle chü, hsü, yü forms above
# Also add variants without the umlaut for broader matching

UMLAUT_VARIANTS = {
    # chü variants
    'chu': 'ju',  # This conflicts with chu -> zhu, so we handle context
    # We keep chü -> ju as the primary mapping

    # hsü variants (when written as hsu)
    'hsu': 'xu',
    'hsuan': 'xuan',
    'hsueh': 'xue',
    'hsun': 'xun',

    # yü variants (when written as yu without umlaut)
    # yu -> you is already in main dict
    'yuan': 'yuan',
    'yueh': 'yue',
    'yun': 'yun',

    # lü variants
    'lu': 'lu',  # Standard mapping
    'lueh': 'lue',

    # nü variants
    'nueh': 'nue',
}

# Merge umlaut variants (only add if not conflicting with existing)
for wg, py in UMLAUT_VARIANTS.items():
    if wg not in WG_TO_PINYIN:
        WG_TO_PINYIN[wg] = py


# Postal romanizations and other common variant spellings
# These are not standard Wade-Giles but appear frequently in older texts
POSTAL_ROMANIZATIONS = {
    # Province names
    'kwangsi': 'guangxi',
    'kwangtung': 'guangdong',
    'fukien': 'fujian',
    'chekiang': 'zhejiang',
    'kiangsi': 'jiangxi',
    'kiangsu': 'jiangsu',
    'shansi': 'shanxi',
    'shensi': 'shaanxi',
    'szechwan': 'sichuan',
    'szechuan': 'sichuan',
    'hopei': 'hebei',
    'hopeh': 'hebei',
    'honan': 'henan',
    'hupei': 'hubei',
    'hupeh': 'hubei',
    'hunan': 'hunan',
    'kansu': 'gansu',
    'kweichow': 'guizhou',
    'yunnan': 'yunnan',
    'anhwei': 'anhui',
    'chihli': 'zhili',
    'fengtien': 'fengtian',
    'manchuria': 'manchuria',  # Keep as is (not Chinese)

    # Major cities
    'peking': 'beijing',
    'peiping': 'beiping',
    'nanking': 'nanjing',
    'canton': 'guangzhou',
    'tientsin': 'tianjin',
    'tsingtao': 'qingdao',
    'chungking': 'chongqing',
    'sian': 'xian',
    'sinkiang': 'xinjiang',
    'tsinghai': 'qinghai',
    'ningsia': 'ningxia',
    'suiyuan': 'suiyuan',

    # Rivers and geographic features
    'yangtze': 'yangzi',
    'yangtse': 'yangzi',

    # Common "kw" forms (Cantonese influence)
    'kwang': 'guang',
    'kwan': 'guan',
    'kwai': 'guai',
    'kwei': 'gui',
}

# Add postal romanizations to main dictionary
for postal, pinyin in POSTAL_ROMANIZATIONS.items():
    if postal not in WG_TO_PINYIN:
        WG_TO_PINYIN[postal] = pinyin


# PDF artifact: "ii" often represents ü (u with umlaut) in converted documents
# This mapping handles syllables where ü appears as "ii"
II_TO_UMLAUT_MAPPINGS = {
    # ch + ii (aspirated, with umlaut) -> q + u
    "ch'ii": 'qu',
    "ch'iian": 'quan',
    "ch'iieh": 'que',
    "ch'iin": 'qun',

    # ch + ii (unaspirated, with umlaut) -> j + u
    'chii': 'ju',
    'chiian': 'juan',
    'chiieh': 'jue',
    'chiin': 'jun',

    # hs + ii -> x + u
    'hsii': 'xu',
    'hsiian': 'xuan',
    'hsiieh': 'xue',
    'hsiin': 'xun',

    # l + ii -> l + ü
    'lii': 'lü',
    'liian': 'luan',
    'liieh': 'lue',

    # n + ii -> n + ü
    'nii': 'nü',
    'niieh': 'nue',

    # y + ii -> y + u
    'yii': 'yu',
    'yiian': 'yuan',
    'yiieh': 'yue',
    'yiin': 'yun',
}

# Add ii mappings to main dictionary
for ii_form, pinyin in II_TO_UMLAUT_MAPPINGS.items():
    if ii_form not in WG_TO_PINYIN:
        WG_TO_PINYIN[ii_form] = pinyin


# Common English words that happen to match Wade-Giles patterns
# These should NOT be converted unless in a clearly Chinese context
# (e.g., hyphenated with other syllables or part of a proper name sequence)
# Only include words where:
# 1. The WG -> PY conversion changes the spelling, AND
# 2. The word is very common in English
ENGLISH_EXCLUSIONS_LOWERCASE = {
    # Very common English function words (lowercase only)
    'to',      # to -> duo (very common English word)
    'no',      # no -> nuo (very common English word)
    'so',      # so -> suo (very common English word)

    # English words where conversion would cause confusion (lowercase only)
    # Note: Capitalized versions may be Chinese proper names, so we only
    # exclude lowercase forms
    'hung',    # hung -> hong (past tense of "hang")
    'sung',    # sung -> song (past tense of "sing")
    'lung',    # lung -> long (body organ)
    'tang',    # tang -> dang (sharp taste/flavor)
    'tan',     # tan -> dan (skin color/tanning)
    'pan',     # pan -> ban (cooking vessel)
    'pen',     # pen -> ben (writing instrument)
    'pin',     # pin -> bin (fastening device)
    'ping',    # ping -> bing (network ping, sound)
    'ting',    # ting -> ding (sound)
}

# Words that should be excluded even when capitalized
# (these are rarely Chinese names)
ENGLISH_EXCLUSIONS_ANY_CASE = {
    'to',
    'no',
    'so',
}

# Syllables that should ONLY be converted in hyphenated Chinese contexts
CONTEXT_SENSITIVE = {
    'i',       # i -> yi (but also English "I")
    'a',       # a -> a (English article)
    'o',       # o -> e (English exclamation)
}


def apply_case(source: str, target: str) -> str:
    """Apply the capitalization pattern from source to target."""
    if not source or not target:
        return target

    result = []
    target_idx = 0

    for i, char in enumerate(source):
        if target_idx >= len(target):
            break

        # Skip apostrophes and hyphens in source when matching
        if char in "''-":
            continue

        if char.isupper():
            result.append(target[target_idx].upper())
        else:
            result.append(target[target_idx].lower())
        target_idx += 1

    # Append any remaining characters from target
    if target_idx < len(target):
        # Determine case from last character of source
        if source and source[-1].isupper():
            result.append(target[target_idx:].upper())
        else:
            result.append(target[target_idx:].lower())

    return ''.join(result)


def normalize_apostrophe(text: str) -> str:
    """Normalize various apostrophe characters to standard apostrophe."""
    # Various apostrophe-like characters
    apostrophes = [''', ''', '`', '´', 'ʼ', 'ʻ', '\u02bc', '\u02bb']
    for apos in apostrophes:
        text = text.replace(apos, "'")
    return text


def normalize_diacritics(text: str) -> str:
    """
    Normalize diacritical marks used in Wade-Giles romanization.

    Some sources use circumflex (^) or other diacritics on vowels.
    For example: Ên, Êrh, Ê (with circumflex) should become En, Erh, E.
    """
    # Mapping of accented characters to their base forms
    diacritic_map = {
        # Uppercase with circumflex
        'Â': 'A', 'Ê': 'E', 'Î': 'I', 'Ô': 'O', 'Û': 'U',
        # Lowercase with circumflex
        'â': 'a', 'ê': 'e', 'î': 'i', 'ô': 'o', 'û': 'u',
        # Uppercase with breve
        'Ă': 'A', 'Ĕ': 'E', 'Ĭ': 'I', 'Ŏ': 'O', 'Ŭ': 'U',
        # Lowercase with breve
        'ă': 'a', 'ĕ': 'e', 'ĭ': 'i', 'ŏ': 'o', 'ŭ': 'u',
        # Uppercase with macron
        'Ā': 'A', 'Ē': 'E', 'Ī': 'I', 'Ō': 'O', 'Ū': 'U',
        # Lowercase with macron
        'ā': 'a', 'ē': 'e', 'ī': 'i', 'ō': 'o', 'ū': 'u',
        # Keep ü as is (it's meaningful in WG)
    }

    for accented, base in diacritic_map.items():
        text = text.replace(accented, base)

    return text


def build_regex_pattern():
    """Build a regex pattern to match all Wade-Giles syllables."""
    # Sort by length (longest first) to ensure longer matches take precedence
    all_syllables = sorted(WG_TO_PINYIN.keys(), key=len, reverse=True)

    # Escape special regex characters and handle apostrophes
    escaped = []
    for syl in all_syllables:
        # Replace apostrophe with pattern matching multiple apostrophe types
        if "'" in syl:
            parts = syl.split("'")
            pattern = "[''`´ʼʻ]?".join(re.escape(p) for p in parts)
            escaped.append(pattern)
        else:
            escaped.append(re.escape(syl))

    # Build pattern with word boundaries
    # Match syllables that are standalone words or connected with hyphens
    pattern = r'\b(' + '|'.join(escaped) + r')\b'

    return pattern


def convert_syllable(match, case_sensitive=True):
    """Convert a matched Wade-Giles syllable to Pinyin."""
    original = match.group(0)
    normalized = normalize_apostrophe(original.lower())

    # Handle ü character (might appear as ü or u with combining umlaut)
    normalized = normalized.replace('ü', 'ü')  # Normalize composed form

    if normalized in WG_TO_PINYIN:
        pinyin = WG_TO_PINYIN[normalized]
        if case_sensitive:
            return apply_case(original, pinyin)
        return pinyin

    return original


class WadeGilesToPinyinConverter:
    """Converter class for Wade-Giles to Pinyin conversion."""

    def __init__(self):
        self.pattern = None
        self._compile_pattern()

    def _compile_pattern(self):
        """Compile the regex pattern for syllable matching."""
        all_syllables = sorted(WG_TO_PINYIN.keys(), key=len, reverse=True)

        patterns = []
        for syl in all_syllables:
            if "'" in syl:
                # Handle apostrophe variants
                parts = syl.split("'")
                pattern = "[''`´ʼʻ']".join(re.escape(p) for p in parts)
                patterns.append(pattern)
            elif 'ü' in syl:
                # Handle umlaut variants
                pattern = re.escape(syl).replace(re.escape('ü'), '[üu]')
                patterns.append(pattern)
            else:
                patterns.append(re.escape(syl))

        # Word boundary pattern
        self.pattern = re.compile(
            r'(?<![a-zA-Z])(' + '|'.join(patterns) + r')(?![a-zA-Z])',
            re.IGNORECASE
        )

    def _convert_single_syllable(self, original: str, aggressive: bool = False,
                                   in_hyphenated_sequence: bool = False) -> str:
        """Convert a single Wade-Giles syllable to Pinyin."""
        normalized = normalize_apostrophe(original.lower())
        normalized = normalized.replace('ü', 'ü')

        # Skip common English words unless in aggressive mode or in a hyphenated sequence
        if not aggressive and not in_hyphenated_sequence:
            if normalized in ENGLISH_EXCLUSIONS_ANY_CASE:
                return original
            if normalized in ENGLISH_EXCLUSIONS_LOWERCASE and original.islower():
                return original

        # Skip context-sensitive syllables unless in hyphenated sequence
        if not aggressive and not in_hyphenated_sequence and normalized in CONTEXT_SENSITIVE:
            return original

        if normalized in WG_TO_PINYIN:
            pinyin = WG_TO_PINYIN[normalized]
            return apply_case(original, pinyin)
        return original

    def _convert_hyphenated_sequence(self, sequence: str, aggressive: bool = False) -> str:
        """
        Convert a hyphenated sequence of Wade-Giles syllables to Pinyin.
        Removes hyphens in the output (e.g., "Tse-tung" -> "Zedong").
        Only first part is capitalized; subsequent parts are lowercase.
        """
        parts = sequence.split('-')
        converted_parts = []
        all_are_wg_syllables = True
        any_converted = False
        is_first_part = True

        for part in parts:
            if not part:
                continue
            # Check if this part matches a WG syllable
            normalized = normalize_apostrophe(part.lower()).replace('ü', 'ü')
            if normalized in WG_TO_PINYIN:
                converted = self._convert_single_syllable(part, aggressive=True,
                                                          in_hyphenated_sequence=True)
                # First part: preserve case; subsequent parts: lowercase
                if is_first_part:
                    converted_parts.append(converted)
                else:
                    converted_parts.append(converted.lower())
                if converted.lower() != part.lower():
                    any_converted = True
            else:
                # Not a WG syllable, keep original
                if is_first_part:
                    converted_parts.append(part)
                else:
                    converted_parts.append(part.lower())
                all_are_wg_syllables = False
            is_first_part = False

        # Remove hyphens if:
        # 1. At least one syllable was actually converted (spelling changed), OR
        # 2. ALL parts are valid WG syllables (likely a Chinese name even if spellings match)
        if any_converted or all_are_wg_syllables:
            return ''.join(converted_parts)
        else:
            return sequence

    def convert_text(self, text: str, aggressive: bool = False) -> str:
        """
        Convert Wade-Giles text to Pinyin.

        Args:
            text: Input text containing Wade-Giles romanization
            aggressive: If True, convert all matches including common English words.
                       If False (default), skip common English words to reduce false positives.

        Returns:
            Text with Wade-Giles converted to Pinyin
        """
        if not text:
            return text

        # First, handle hyphenated sequences (e.g., "Tse-tung" -> "Zedong")
        # Match sequences of syllables connected by hyphens
        hyphen_pattern = re.compile(
            r'\b([a-zA-ZüÜ\'\'`´ʼʻ]+(?:-[a-zA-ZüÜ\'\'`´ʼʻ]+)+)\b'
        )

        def hyphen_replacer(match):
            return self._convert_hyphenated_sequence(match.group(1), aggressive)

        text = hyphen_pattern.sub(hyphen_replacer, text)

        # Then handle standalone syllables
        def replacer(match):
            original = match.group(1)
            return self._convert_single_syllable(original, aggressive, in_hyphenated_sequence=False)

        return self.pattern.sub(replacer, text)

    def _process_textboxes_in_xml(self, xml_content: bytes, aggressive: bool = False) -> bytes:
        """
        Process text boxes in the document XML using efficient regex-based replacement.

        This approach is much faster and more memory-efficient than DOM parsing
        for large documents. It finds all <w:t>...</w:t> elements within
        <w:txbxContent> sections and applies conversion to their text content.

        Args:
            xml_content: The raw XML content of document.xml
            aggressive: If True, convert all matches including common English words

        Returns:
            Modified XML content as bytes
        """
        # Decode to string for processing
        xml_str = xml_content.decode('utf-8')

        conversion_count = 0

        # Find all txbxContent sections and process text within them
        # Pattern to find text box content sections
        txbx_pattern = re.compile(
            r'(<w:txbxContent[^>]*>)(.*?)(</w:txbxContent>)',
            re.DOTALL
        )

        def process_textbox_content(match):
            nonlocal conversion_count
            start_tag = match.group(1)
            content = match.group(2)
            end_tag = match.group(3)

            # Within textbox content, find and convert text in <w:t> elements
            # Pattern for <w:t>text</w:t> or <w:t xml:space="preserve">text</w:t>
            text_pattern = re.compile(r'(<w:t[^>]*>)([^<]*)(</w:t>)')

            def convert_text_element(text_match):
                nonlocal conversion_count
                open_tag = text_match.group(1)
                text = text_match.group(2)
                close_tag = text_match.group(3)

                if text:
                    converted = self.convert_text(text, aggressive=aggressive)
                    if converted != text:
                        conversion_count += 1
                        return open_tag + converted + close_tag
                return text_match.group(0)

            processed_content = text_pattern.sub(convert_text_element, content)
            return start_tag + processed_content + end_tag

        modified_xml = txbx_pattern.sub(process_textbox_content, xml_str)

        if conversion_count > 0:
            print(f"  Converted {conversion_count} text elements in text boxes")

        return modified_xml.encode('utf-8')

    def _process_all_text_in_xml(self, xml_content: bytes, aggressive: bool = False) -> bytes:
        """
        Process ALL text in the document XML using regex-based replacement.

        This method processes all <w:t> elements in the document, not just those
        in text boxes. Used as a fallback when python-docx fails to load large documents.

        Args:
            xml_content: The raw XML content of document.xml
            aggressive: If True, convert all matches including common English words

        Returns:
            Modified XML content as bytes
        """
        # Decode to string for processing
        xml_str = xml_content.decode('utf-8')

        conversion_count = 0

        # Pattern for all <w:t>text</w:t> elements
        text_pattern = re.compile(r'(<w:t[^>]*>)([^<]*)(</w:t>)')

        def convert_text_element(text_match):
            nonlocal conversion_count
            open_tag = text_match.group(1)
            text = text_match.group(2)
            close_tag = text_match.group(3)

            if text:
                converted = self.convert_text(text, aggressive=aggressive)
                if converted != text:
                    conversion_count += 1
                    return open_tag + converted + close_tag
            return text_match.group(0)

        modified_xml = text_pattern.sub(convert_text_element, xml_str)

        print(f"  Converted {conversion_count} text elements")

        return modified_xml.encode('utf-8')

    def _convert_docx_via_xml(self, input_path: Path, output_path: Path,
                               aggressive: bool = False) -> str:
        """
        Convert a .docx file using direct XML manipulation only.

        This is used as a fallback when python-docx fails to load the document
        (e.g., due to XML parser limits with very large documents).

        Args:
            input_path: Path to input .docx file
            output_path: Path for output file
            aggressive: If True, convert all matches including common English words

        Returns:
            Path to the output file
        """
        # Copy the input file to output first
        shutil.copy2(input_path, output_path)

        # Process the XML directly
        with zipfile.ZipFile(output_path, 'r') as zf:
            xml_content = zf.read('word/document.xml')

            # Process ALL text in the document
            modified_xml = self._process_all_text_in_xml(xml_content, aggressive)

            # Read all other files
            other_files = {}
            for item in zf.namelist():
                if item != 'word/document.xml':
                    other_files[item] = zf.read(item)

        # Write the modified docx
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr('word/document.xml', modified_xml)
            for name, content in other_files.items():
                zf.writestr(name, content)

        return str(output_path)

    def convert_docx(self, input_path: str, output_path: str = None,
                      aggressive: bool = False) -> str:
        """
        Convert Wade-Giles in a .docx file to Pinyin.

        This method processes:
        - Regular paragraphs
        - Tables
        - Headers and footers
        - Text boxes (via direct XML manipulation)

        For very large documents that exceed XML parser limits, it automatically
        falls back to direct XML processing.

        Args:
            input_path: Path to input .docx file
            output_path: Path for output file (optional)
            aggressive: If True, convert all matches including common English words

        Returns:
            Path to the output file
        """
        input_path = Path(input_path)

        if output_path is None:
            output_path = input_path.parent / f"{input_path.stem}_pinyin{input_path.suffix}"
        else:
            output_path = Path(output_path)

        # Try to process with python-docx first
        try:
            print("Processing regular content (paragraphs, tables, headers/footers)...")
            doc = Document(str(input_path))

            # Process paragraphs
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.text:
                        run.text = self.convert_text(run.text, aggressive=aggressive)

            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text:
                                    run.text = self.convert_text(run.text, aggressive=aggressive)

            # Process headers and footers
            for section in doc.sections:
                header = section.header
                for para in header.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = self.convert_text(run.text, aggressive=aggressive)

                footer = section.footer
                for para in footer.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = self.convert_text(run.text, aggressive=aggressive)

            # Save intermediate result
            doc.save(str(output_path))

            # Now process text boxes by directly modifying the XML
            print("Processing text boxes...")
            self._process_textboxes_in_docx(output_path, aggressive)

        except Exception as e:
            # Check if this is an XML parser error (common with large PDF-converted docs)
            error_msg = str(e)
            if 'XMLSyntaxError' in type(e).__name__ or 'AttValue' in error_msg or 'lxml' in error_msg:
                print(f"Document too large for standard processing. Using direct XML mode...")
                print("Processing all text via XML manipulation...")
                return self._convert_docx_via_xml(input_path, output_path, aggressive)
            else:
                # Re-raise other errors
                raise

        return str(output_path)

    def _process_textboxes_in_docx(self, docx_path: Path, aggressive: bool = False):
        """
        Process text boxes in a .docx file by modifying its XML directly.

        Args:
            docx_path: Path to the .docx file to modify in place
            aggressive: If True, convert all matches including common English words
        """
        docx_path = Path(docx_path)

        # Read the docx as a zip file
        with zipfile.ZipFile(docx_path, 'r') as zf:
            # Read document.xml
            xml_content = zf.read('word/document.xml')

            # Check if there are any text boxes
            if b'txbxContent' not in xml_content:
                print("  No text boxes found in document")
                return

            # Process the XML
            modified_xml = self._process_textboxes_in_xml(xml_content, aggressive)

            # Read all other files
            other_files = {}
            for item in zf.namelist():
                if item != 'word/document.xml':
                    other_files[item] = zf.read(item)

        # Write the modified docx
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            # Write modified document.xml
            zf.writestr('word/document.xml', modified_xml)

            # Write all other files unchanged
            for name, content in other_files.items():
                zf.writestr(name, content)

    def _convert_pdf_word(self, word_clean: str, wg_lookup: dict) -> str:
        """
        Convert a single word (possibly hyphenated) from Wade-Giles to Pinyin.

        Handles:
        - Diacritic normalization (Ê -> E)
        - Apostrophe normalization
        - Hyphenated names (split, convert each part, rejoin without hyphens)

        Args:
            word_clean: The word to convert (punctuation already stripped)
            wg_lookup: Dictionary mapping lowercase WG terms to pinyin

        Returns:
            Converted pinyin string, or None if no conversion applicable
        """
        # Normalize diacritics and apostrophes
        normalized = normalize_diacritics(word_clean)
        normalized = normalize_apostrophe(normalized)

        # Check if it's a hyphenated name
        if '-' in normalized:
            parts = normalized.split('-')
            converted_parts = []
            any_converted = False
            all_are_wg = True
            is_first_part = True

            for part in parts:
                part_lower = part.lower()
                # Check if it's a valid WG syllable (in full dictionary)
                is_wg_syllable = part_lower in WG_TO_PINYIN
                if not is_wg_syllable:
                    all_are_wg = False

                if part_lower in wg_lookup:
                    # This syllable changes when converted
                    pinyin = wg_lookup[part_lower]
                    if is_first_part:
                        converted_parts.append(apply_case(part, pinyin))
                    else:
                        converted_parts.append(pinyin.lower())
                    any_converted = True
                elif is_wg_syllable:
                    # Valid WG syllable but doesn't change (e.g., 'en' -> 'en')
                    pinyin = WG_TO_PINYIN[part_lower]
                    if is_first_part:
                        converted_parts.append(apply_case(part, pinyin))
                    else:
                        converted_parts.append(pinyin.lower())
                else:
                    # Not a WG syllable, keep original
                    if is_first_part:
                        converted_parts.append(part)
                    else:
                        converted_parts.append(part.lower())
                is_first_part = False

            # Return joined form if any part was converted OR all parts are valid WG
            if any_converted or all_are_wg:
                return ''.join(converted_parts)
            return None

        # Single word lookup
        word_lower = normalized.lower()
        if word_lower in wg_lookup:
            pinyin = wg_lookup[word_lower]
            return apply_case(normalized, pinyin)

        return None

    def _is_likely_proper_noun(self, word_clean: str, word_lower: str,
                                is_sentence_start: bool, wg_lookup: dict) -> bool:
        """
        Determine if a word is likely a proper noun that should be converted.

        Args:
            word_clean: The cleaned word (punctuation stripped)
            word_lower: Lowercase normalized version
            is_sentence_start: Whether this word is at the start of a sentence
            wg_lookup: The WG lookup dictionary

        Returns:
            True if the word should be converted
        """
        # Must be capitalized
        if not word_clean[0].isupper():
            return False

        # Normalize for checking
        normalized = normalize_diacritics(word_clean)
        normalized = normalize_apostrophe(normalized)
        normalized_lower = normalized.lower()

        # If it contains apostrophe or hyphen, it's clearly WG - always convert
        if "'" in normalized or "-" in normalized:
            return True

        # Check if it's a known postal romanization (always proper nouns)
        if normalized_lower in POSTAL_ROMANIZATIONS:
            return True

        # At sentence start, be more conservative
        if is_sentence_start:
            # Only convert at sentence start if it's multi-character and distinctive
            # (longer than 4 chars suggests a distinctive WG term)
            if len(normalized) > 4:
                return True
            # Skip short words at sentence start (could be English)
            return False

        # Not at sentence start - if capitalized and in dictionary, convert it
        return True

    def convert_pdf(self, input_path: str, output_path: str = None,
                    aggressive: bool = False) -> str:
        """
        Convert Wade-Giles in a PDF file to Pinyin.

        Uses search_for to find exact text matches and get precise bounding boxes.
        Handles diacritics (Ê, ê) and hyphenated names (Tse-tung -> Zedong).
        Only converts proper nouns (capitalized words).
        Uses Times New Roman (serif) font at 10.5pt.

        WARNING: PDF conversion has limitations:
        - Text positioning may be slightly off after replacement
        - Different-length replacements may cause visual artifacts
        - Font matching is approximate

        Args:
            input_path: Path to input PDF file
            output_path: Path for output file (optional)
            aggressive: If True, convert all matches including common English words

        Returns:
            Path to the output file
        """
        if not PDF_SUPPORT:
            raise ImportError(
                "PDF support requires PyMuPDF. Install with: pip install pymupdf"
            )

        input_path = Path(input_path)

        if output_path is None:
            output_path = input_path.parent / f"{input_path.stem}_pinyin{input_path.suffix}"
        else:
            output_path = Path(output_path)

        # Build search terms: list of (search_term, replacement) tuples
        # Only search for capitalized terms (proper nouns)
        search_terms = self._build_pdf_search_list(aggressive)

        print(f"Opening PDF: {input_path}")
        print(f"Loaded {len(search_terms)} search terms for conversion")
        doc = fitz.open(str(input_path))
        total_pages = len(doc)
        total_replacements = 0

        # Fixed font size for consistent appearance
        FONT_SIZE = 10.5

        print(f"Processing {total_pages} pages...")

        for page_num in range(total_pages):
            page = doc[page_num]
            page_replacements = 0

            # Collect all replacements to make on this page
            # Each item is (rect, replacement_text, original_word)
            replacements_to_make = []

            # Get all words on this page for exact matching
            page_words = page.get_text("words")
            # Build a dict of normalized word -> list of (rect, original_word)
            word_positions = {}
            for x0, y0, x1, y1, word, *_ in page_words:
                # Normalize the word for lookup
                word_clean = word.strip('.,;:!?()[]"\'')
                if not word_clean or len(word_clean) < 2:
                    continue
                # Normalize diacritics and apostrophes
                normalized = normalize_diacritics(word_clean)
                normalized = normalize_apostrophe(normalized)
                key = normalized
                if key not in word_positions:
                    word_positions[key] = []
                word_positions[key].append((fitz.Rect(x0, y0, x1, y1), word))

            # Build a lookup from search terms for single-word matching
            single_word_lookup = {}
            for search_term, replacement in search_terms:
                # Normalize the search term
                search_normalized = normalize_diacritics(search_term)
                search_normalized = normalize_apostrophe(search_normalized)
                single_word_lookup[search_normalized] = replacement

            # Also need lookup for individual syllables (for hyphenated words)
            wg_lookup = self._build_pdf_lookup()

            # Process each word position
            for word_key, positions in word_positions.items():
                for rect, original_word in positions:
                    # Only convert capitalized words (proper nouns)
                    word_clean = original_word.strip('.,;:!?()[]"\'')
                    if not word_clean or not word_clean[0].isupper():
                        continue

                    # Normalize for lookup
                    normalized = normalize_diacritics(word_clean)
                    normalized = normalize_apostrophe(normalized)

                    replacement = None

                    # First try exact match in search terms
                    if normalized in single_word_lookup:
                        replacement = single_word_lookup[normalized]
                    # Then try hyphenated word handling
                    elif '-' in normalized:
                        parts = normalized.split('-')
                        converted_parts = []
                        any_converted = False
                        is_first_part = True
                        for part in parts:
                            part_lower = part.lower()
                            if part_lower in wg_lookup:
                                pinyin = wg_lookup[part_lower]
                                # First part: preserve case; subsequent parts: lowercase
                                if is_first_part:
                                    converted_parts.append(apply_case(part, pinyin))
                                else:
                                    converted_parts.append(pinyin.lower())
                                any_converted = True
                            else:
                                if is_first_part:
                                    converted_parts.append(part)
                                else:
                                    converted_parts.append(part.lower())
                            is_first_part = False
                        if any_converted:
                            replacement = ''.join(converted_parts)
                    # Finally try single syllable lookup
                    else:
                        normalized_lower = normalized.lower()
                        if normalized_lower in wg_lookup:
                            replacement = apply_case(normalized, wg_lookup[normalized_lower])

                    if replacement is None:
                        continue

                    # Build replacement preserving punctuation
                    leading = ""
                    trailing = ""
                    for c in original_word:
                        if c in '.,;:!?()[]"\' ':
                            leading += c
                        else:
                            break
                    for c in reversed(original_word):
                        if c in '.,;:!?()[]"\' ':
                            trailing = c + trailing
                        else:
                            break
                    full_replacement = leading + replacement + trailing

                    replacements_to_make.append((rect, full_replacement, original_word))

            # Remove duplicates (same rect matched by multiple search terms)
            seen_rects = set()
            unique_replacements = []
            for item in replacements_to_make:
                rect = item[0]
                rect_key = (round(rect.x0, 1), round(rect.y0, 1), round(rect.x1, 1), round(rect.y1, 1))
                if rect_key not in seen_rects:
                    seen_rects.add(rect_key)
                    unique_replacements.append(item)

            # Sort by position (top to bottom, left to right)
            unique_replacements.sort(key=lambda x: (x[0].y0, x[0].x0))

            # Apply all replacements using overlay method (draw white rect + new text)
            # This avoids the redaction issues that affect surrounding text
            shape = page.new_shape()
            for rect, replacement, _ in unique_replacements:
                # Draw white rectangle to cover original text
                shape.draw_rect(rect)
                shape.finish(color=(1, 1, 1), fill=(1, 1, 1))
                page_replacements += 1
            shape.commit()

            # Insert replacement text
            for rect, replacement, _ in unique_replacements:
                # Calculate insertion point for baseline
                insert_point = fitz.Point(rect.x0, rect.y1 - 2)
                page.insert_text(
                    insert_point,
                    replacement,
                    fontsize=FONT_SIZE,
                    fontname="times-roman",
                    color=(0, 0, 0)
                )

            total_replacements += page_replacements

            # Progress indicator every 50 pages
            if (page_num + 1) % 50 == 0 or page_num == total_pages - 1:
                print(f"  Processed {page_num + 1}/{total_pages} pages...")

        print(f"Saving converted PDF to: {output_path}")
        doc.save(str(output_path), garbage=4, deflate=True)
        doc.close()

        print(f"Total replacements: {total_replacements}")
        return str(output_path)

    def _build_pdf_lookup(self) -> dict:
        """
        Build a lookup dictionary for PDF word-by-word conversion.

        Returns a dictionary mapping lowercase WG terms to their pinyin equivalents.
        Excludes common English words to prevent false positives.

        Returns:
            Dictionary mapping lowercase WG terms to pinyin
        """
        lookup = {}

        # Add all WG_TO_PINYIN entries
        for wg, pinyin in WG_TO_PINYIN.items():
            # Skip if identical
            if wg == pinyin:
                continue
            # Skip common English words
            if wg in ENGLISH_EXCLUSIONS_LOWERCASE or wg in ENGLISH_EXCLUSIONS_ANY_CASE:
                continue
            # Skip very short terms
            if len(wg) < 2:
                continue
            # Normalize and add
            lookup[wg.lower()] = pinyin

        # Add hyphenated name components that should be joined
        # These are handled specially - the hyphenated form maps to joined pinyin
        hyphenated = {
            'tse-tung': 'zedong',
            'en-lai': 'enlai',
            "hsiao-p'ing": 'xiaoping',
            'kai-shek': 'jieshi',
            'yat-sen': 'yixian',
            'chung-shan': 'zhongshan',
        }
        lookup.update(hyphenated)

        return lookup

    def _build_pdf_search_list(self, aggressive: bool = False) -> list:
        """
        Build a list of (search_term, replacement) tuples for PDF conversion.

        Only includes capitalized terms (proper nouns) to avoid false positives.
        Sorted by length (longest first) so longer matches take precedence.
        Includes diacritic variants (Ê, ê) for comprehensive matching.

        Args:
            aggressive: If True, include lowercase terms and English words

        Returns:
            List of (search_term, replacement) tuples, sorted by length descending
        """
        terms = []
        seen = set()  # Avoid duplicates

        # Build lookup for conversion
        wg_lookup = self._build_pdf_lookup()

        # Process all WG terms, creating capitalized search terms
        for wg, pinyin in WG_TO_PINYIN.items():
            # Skip if identical
            if wg == pinyin:
                continue
            # Skip very short terms
            if len(wg) < 2:
                continue
            # Skip common English words unless aggressive
            if not aggressive:
                if wg in ENGLISH_EXCLUSIONS_LOWERCASE or wg in ENGLISH_EXCLUSIONS_ANY_CASE:
                    continue

            # Only add capitalized version for proper nouns
            cap_wg = wg.capitalize()
            cap_pinyin = pinyin.capitalize()
            if cap_wg not in seen:
                terms.append((cap_wg, cap_pinyin))
                seen.add(cap_wg)

            # Add diacritic variants (Ê for E, etc.)
            # Map base vowels to their diacritic forms
            diacritic_map = {'E': 'Ê', 'e': 'ê', 'A': 'Â', 'a': 'â',
                             'O': 'Ô', 'o': 'ô', 'U': 'Û', 'u': 'û', 'I': 'Î', 'i': 'î'}
            for base, diac in diacritic_map.items():
                if base in cap_wg:
                    variant = cap_wg.replace(base, diac)
                    if variant not in seen:
                        terms.append((variant, cap_pinyin))
                        seen.add(variant)

        # Add apostrophe variants
        apostrophe_chars = ["'", "'", "'", "`", "´"]
        new_terms = []
        for search, repl in terms:
            if "'" in search:
                for apos in apostrophe_chars:
                    variant = search.replace("'", apos)
                    if variant not in seen:
                        new_terms.append((variant, repl))
                        seen.add(variant)
        terms.extend(new_terms)

        # Add common hyphenated names (second part lowercase per Pinyin convention)
        hyphenated = [
            ('Ên-Ssu', 'Ensi'), ('Ên-Fu', 'Enfu'),
            ('Êrh-Yuan', 'Eryuan'), ('Êrh-Chou', 'Erzhou'),
            ('Tse-tung', 'Zedong'), ('Tse-Tung', 'Zedong'),
            ('En-lai', 'Enlai'), ('En-Lai', 'Enlai'),
            ("Ch'ê-Chiang", 'Chejiang'), ("Ch'ê-Men", 'Chemen'),
            ("Ch'ên-Chiang", 'Chenjiang'), ("Ch'ên-Pu", 'Chenbu'),
            ("Ch'êng-Kung", 'Chenggong'), ("Ch'êng-Ho", 'Chenghe'),
            ("Hsiao-p'ing", 'Xiaoping'), ("Hsiao-P'ing", 'Xiaoping'),
            ('Kai-shek', 'Jieshi'), ('Kai-Shek', 'Jieshi'),
            ('Yat-sen', 'Yixian'), ('Yat-Sen', 'Yixian'),
            ('Chung-shan', 'Zhongshan'), ('Chung-Shan', 'Zhongshan'),
        ]
        for search, repl in hyphenated:
            if search not in seen:
                terms.append((search, repl))
                seen.add(search)
            # Also add apostrophe variants
            for apos in apostrophe_chars:
                variant = search.replace("'", apos)
                if variant not in seen:
                    terms.append((variant, repl))
                    seen.add(variant)

        # Sort by length (longest first) so longer matches take precedence
        terms.sort(key=lambda x: len(x[0]), reverse=True)

        return terms

    def _build_pdf_search_terms(self, aggressive: bool = False) -> dict:
        """
        Build a dictionary of search terms for PDF conversion.

        Uses the FULL WG_TO_PINYIN dictionary to convert all Wade-Giles terms.
        Terms are sorted by length (longest first) and include case variants.

        Args:
            aggressive: If True, include terms that might be English words

        Returns:
            Dictionary mapping search terms to their Pinyin replacements
        """
        terms = {}

        # Process ALL entries from WG_TO_PINYIN
        for wg, pinyin in WG_TO_PINYIN.items():
            # Skip if WG and Pinyin are identical (no change needed)
            if wg == pinyin:
                continue

            # Skip very short terms that are likely false positives
            if len(wg) < 2:
                continue

            # Skip common English words unless aggressive mode
            if not aggressive:
                if wg in ENGLISH_EXCLUSIONS_ANY_CASE:
                    continue
                if wg in ENGLISH_EXCLUSIONS_LOWERCASE:
                    # Only add capitalized version
                    cap_wg = wg.capitalize()
                    cap_pinyin = pinyin.capitalize()
                    terms[cap_wg] = cap_pinyin
                    continue

            # Add capitalized version (most common in proper names)
            cap_wg = wg.capitalize()
            cap_pinyin = pinyin.capitalize()
            terms[cap_wg] = cap_pinyin

            # Add lowercase version for words in running text
            terms[wg] = pinyin

            # Add UPPERCASE version for titles/headers
            terms[wg.upper()] = pinyin.upper()

        # Handle apostrophe variants for aspirated consonants
        # PDFs may use different apostrophe characters
        apostrophe_variants = ["'", "'", "'", "`", "´"]
        terms_with_apostrophes = {}

        for wg, pinyin in list(terms.items()):
            if "'" in wg:
                # Add variants with different apostrophe characters
                for apos in apostrophe_variants:
                    variant = wg.replace("'", apos)
                    if variant not in terms:
                        terms_with_apostrophes[variant] = pinyin

        terms.update(terms_with_apostrophes)

        # Add common hyphenated names (convert and remove hyphen)
        hyphenated_names = {
            'Tse-tung': 'Zedong',
            'Tse-Tung': 'Zedong',
            'TSE-TUNG': 'ZEDONG',
            'En-lai': 'Enlai',
            'En-Lai': 'Enlai',
            'EN-LAI': 'ENLAI',
            "Hsiao-p'ing": 'Xiaoping',
            "Hsiao-P'ing": 'Xiaoping',
            'Kai-shek': 'Jieshi',
            'Kai-Shek': 'Jieshi',
            'Yat-sen': 'Yixian',
            'Yat-Sen': 'Yixian',
            'Chung-shan': 'Zhongshan',
            'Chung-Shan': 'Zhongshan',
            'Chiang-nan': 'Jiangnan',
            'Hua-pei': 'Huabei',
            'Hua-nan': 'Huanan',
        }
        terms.update(hyphenated_names)

        # Add Taoism/Daoism variants
        tao_terms = {
            'Taoism': 'Daoism',
            'TAOISM': 'DAOISM',
            'taoism': 'daoism',
            'Taoist': 'Daoist',
            'TAOIST': 'DAOIST',
            'taoist': 'daoist',
            'Taoists': 'Daoists',
            'taoists': 'daoists',
        }
        terms.update(tao_terms)

        return terms


def main():
    """Main entry point for command-line usage."""
    import argparse

    parser = argparse.ArgumentParser(
        description='Convert Wade-Giles romanization to Pinyin in .docx or .pdf files.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python wg_to_pinyin.py document.docx
  python wg_to_pinyin.py input.docx output.docx
  python wg_to_pinyin.py input.pdf output.pdf
  python wg_to_pinyin.py input.docx -o output.docx --aggressive

Notes:
  - Supports both .docx and .pdf files (PDF requires: pip install pymupdf)
  - By default, common English words (to, no, so, lung, tang, etc.) are NOT
    converted to avoid false positives.
  - Capitalized versions of ambiguous words ARE converted, as they may be
    Chinese proper names (e.g., "Sung Dynasty" -> "Song Dynasty").
  - Use --aggressive to convert ALL matches including common English words.

PDF Limitations:
  - Text positioning may be slightly off after replacement
  - Different-length replacements may cause visual artifacts
  - Font matching is approximate
        """
    )

    parser.add_argument('input', help='Input .docx or .pdf file')
    parser.add_argument('output', nargs='?', help='Output file (default: input_pinyin.ext)')
    parser.add_argument('-o', '--output-file', dest='output_file',
                        help='Output file (alternative to positional argument)')
    parser.add_argument('-a', '--aggressive', action='store_true',
                        help='Convert all matches, including common English words')

    args = parser.parse_args()

    input_file = args.input
    output_file = args.output or args.output_file
    input_path = Path(input_file)

    if not input_path.exists():
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)

    converter = WadeGilesToPinyinConverter()

    # Determine file type and convert accordingly
    suffix = input_path.suffix.lower()

    if suffix == '.pdf':
        if not PDF_SUPPORT:
            print("Error: PDF support requires PyMuPDF. Install with: pip install pymupdf")
            sys.exit(1)
        output_path = converter.convert_pdf(input_file, output_file, aggressive=args.aggressive)
    elif suffix in ['.docx', '.doc']:
        output_path = converter.convert_docx(input_file, output_file, aggressive=args.aggressive)
    else:
        print(f"Error: Unsupported file type '{suffix}'. Supported: .docx, .pdf")
        sys.exit(1)

    print(f"Conversion complete. Output saved to: {output_path}")


if __name__ == '__main__':
    main()
